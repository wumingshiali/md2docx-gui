"""
任务列表转换器模块
"""
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

from .base import ElementConverter
from .list import ListConverter


class TaskListConverter(ElementConverter):
    """任务列表转换器，处理Markdown中的任务列表（TODO列表）"""

    def __init__(self, base_converter=None):
        """初始化任务列表转换器
        
        Args:
            base_converter: 基础转换器实例
        """
        super().__init__(base_converter)
        self.debug = False
        self.list_converter = None
        if base_converter:
            self.debug = base_converter.debug
            # 获取列表转换器，用于处理基本列表结构
            if 'list' in base_converter.converters:
                self.list_converter = base_converter.converters['list']
            else:
                self.list_converter = ListConverter(base_converter)

    def convert(self, tokens):
        """转换任务列表token为DOCX带符号的列表
        
        Args:
            tokens: 任务列表token元组 (list_token, content_token)
            
        Returns:
            docx.paragraph: 创建的段落对象
        """
        if not self.document:
            raise ValueError("Document not set for TaskListConverter")
        
        if self.debug:
            print(f"处理任务列表: {tokens}")
        
        # 解析token
        list_token, content_token = tokens
        
        # 检查是否为任务列表项
        is_checked = False
        task_text = ""
        
        # 清理任务标记的正则表达式
        task_pattern = re.compile(r'^\s*\[([ x])\]\s*')
        
        if hasattr(content_token, 'content'):
            content = content_token.content
            # 使用正则表达式检查和移除任务标记
            try:
                match = task_pattern.match(content)
                if match:
                    is_checked = match.group(1) == 'x'
                    task_text = task_pattern.sub('', content)
                else:
                    task_text = content
            except TypeError:
                # 处理 content 不是字符串的情况
                task_text = str(content) if content is not None else ""
        elif hasattr(content_token, 'children'):
            for child in content_token.children:
                if hasattr(child, 'type') and child.type == 'checkbox_input':
                    is_checked = child.attrs.get('checked', False) if hasattr(child, 'attrs') else False
                elif hasattr(child, 'content'):
                    # 使用正则表达式移除内容中的任务标记
                    try:
                        child_content = child.content
                        match = task_pattern.match(child_content)
                        if match:
                            child_content = task_pattern.sub('', child_content)
                        task_text += child_content
                    except TypeError:
                        # 处理 child.content 不是字符串的情况
                        task_text += str(child.content) if child.content is not None else ""
        
        # 使用符号替代复选框
        checkbox_symbol = "√ " if is_checked else "× "
        
        # 将符号添加到任务文本前面
        task_text_with_symbol = checkbox_symbol + task_text
        
        # 使用列表转换器创建段落
        paragraph = None
        if self.list_converter:
            try:
                # 创建一个新的内容token，只包含任务文本（不包含符号）
                new_content_token = type('InlineToken', (), {
                    'type': 'inline',
                    'content': task_text,  # 不包含符号，以便列表转换器正常处理
                    'children': []
                })
                
                # 使用列表转换器创建段落
                paragraph = self.list_converter.convert((list_token, new_content_token))
                
                # 清空段落内容
                if paragraph.runs:
                    # 直接修改第一个run的文本，而不是清空后重新添加
                    paragraph.runs[0].text = task_text_with_symbol
                    
                    # 删除其他所有runs（如果有的话）
                    for i in range(len(paragraph.runs) - 1, 0, -1):
                        paragraph._p.remove(paragraph.runs[i]._r)
                else:
                    paragraph.add_run(task_text_with_symbol)
                
                return paragraph
            except Exception as e:
                if self.debug:
                    print(f"使用列表转换器创建段落时出错: {e}")
        
        # 如果列表转换器失败或不存在，创建一个简单的段落
        paragraph = self.document.add_paragraph()
        paragraph.add_run(task_text_with_symbol)
        return paragraph

    def _add_checkbox(self, paragraph, is_checked=False):
        """向段落添加复选框
        
        Args:
            paragraph: 段落对象
            is_checked: 是否勾选
        """
        # 检查段落是否为None
        if paragraph is None:
            if self.debug:
                print("警告: 尝试向None段落添加复选框")
            return
            
        # 获取段落的第一个run
        if not paragraph.runs:
            run = paragraph.add_run()
        else:
            run = paragraph.runs[0]
            
        # 添加复选框符号
        if is_checked:
            run.text = "√ " + run.text
        else:
            run.text = "× " + run.text 