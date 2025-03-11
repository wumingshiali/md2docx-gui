"""
引用块转换器模块，处理引用块的转换
"""
from typing import Any, Optional, Tuple
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import ElementConverter


class BlockquoteConverter(ElementConverter):
    """处理引用块的转换器"""
    
    def __init__(self, base_converter=None):
        super().__init__(base_converter)
    
    def convert(self, tokens: Tuple[Any, Any]) -> None:
        """转换引用块元素
        
        Args:
            tokens: (开始标记, 内容标记) 的元组
        """
        if not self.document:
            raise ValueError("Document not set")
            
        quote_token, content_token = tokens
        
        # 获取引用块层级
        level = len(quote_token.markup) if hasattr(quote_token, 'markup') else 1
        
        # 创建或获取引用块样式
        style_name = "Quote" if level == 1 else f"Quote{level}"
        self._ensure_quote_style(style_name, level)
        
        # 创建新段落
        paragraph = self.document.add_paragraph()
        paragraph.style = self.document.styles[style_name]
        
        # 处理空引用块
        if not content_token:
            paragraph.add_run("")
            return
        
        # 处理引用块内容
        current_text = ""
        current_style = {"bold": False, "italic": False, "strike": False}
        
        for child in content_token.children:
            if child.type == 'text':
                # 处理多行文本中的空格
                text = child.content.replace('\n', ' ')
                if text.endswith(' '):
                    text = text[:-1]
                current_text += text
            elif child.type == 'strong_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["bold"] = True
            elif child.type == 'strong_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["bold"] = False
            elif child.type == 'em_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["italic"] = True
            elif child.type == 'em_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["italic"] = False
            elif child.type == 'softbreak':
                current_text += " "
        
        # 添加剩余的文本
        if current_text:
            self._add_text_with_style(paragraph, current_text, current_style)
    
    def _add_text_with_style(self, paragraph, text: str, style: dict) -> None:
        """添加带样式的文本
        
        Args:
            paragraph: 段落对象
            text: 要添加的文本
            style: 样式配置
        """
        run = paragraph.add_run(text)
        run.bold = style["bold"]
        run.italic = style["italic"]
    
    def _ensure_quote_style(self, style_name: str, level: int) -> None:
        """确保引用块样式存在
        
        Args:
            style_name: 样式名称
            level: 引用块层级
        """
        if style_name not in self.document.styles:
            style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            # 设置基本样式
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(102, 102, 102)  # 灰色
            # 根据层级设置左缩进
            style.paragraph_format.left_indent = Pt(30 * level)
            # 设置段落间距
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            # 设置对齐方式
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT 