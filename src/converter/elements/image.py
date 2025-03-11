"""
图片转换器模块
"""
import os
import re
import requests
from io import BytesIO
from typing import Any, Dict, Optional, Tuple
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import ElementConverter


class ImageConverter(ElementConverter):
    """图片转换器，处理各种类型的图片"""

    def __init__(self, base_converter=None):
        super().__init__(base_converter)
        self.document = None
        # 图片缓存，避免重复下载
        self._image_cache = {}

    def convert(self, tokens: Tuple[Any, Any]) -> None:
        """转换图片元素
        
        Args:
            tokens: (开始标记, 内容标记) 的元组
        """
        if not self.document:
            raise ValueError("Document not set")
            
        token, content_token = tokens
        
        # 获取图片信息
        if not hasattr(token, 'attrs') or not token.attrs:
            return
        
        # 获取图片URL和标题
        src = token.attrs.get('src', '')
        title = token.attrs.get('title', '')
        
        # 获取alt文本
        alt = ""
        if hasattr(token, 'content'):
            alt = token.content
        elif hasattr(content_token, 'content'):
            alt = content_token.content
        
        # 调试信息
        debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
        if debug:
            print(f"处理图片: src={src}, alt={alt}, title={title}")
        
        # 解析尺寸信息（如果有）
        width, height = self._parse_size(alt)
        if width and height and debug:
            print(f"图片尺寸: {width}x{height}")
        
        # 创建段落并设置居中对齐
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加图片
        try:
            # 获取图片数据
            image_data = self._get_image_data(src)
            if not image_data:
                if debug:
                    print(f"无法获取图片数据: {src}")
                return
            
            # 添加图片到文档
            if width and height:
                # 使用指定尺寸
                run = paragraph.add_run()
                run.add_picture(image_data, width=Pt(width), height=Pt(height))
            else:
                # 使用默认尺寸
                run = paragraph.add_run()
                run.add_picture(image_data)
            
            # 添加图片标题（如果有）
            if title:
                caption_paragraph = self.document.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_paragraph.add_run(title)
                caption_run.italic = True
                caption_run.font.size = Pt(10)
            
            if debug:
                print(f"图片添加成功: {src}")
                
        except Exception as e:
            if debug:
                print(f"添加图片失败: {str(e)}")
    
    def convert_in_paragraph(self, paragraph, token, style=None):
        """在段落中转换图片
        
        Args:
            paragraph: 段落对象
            token: 图片标记
            style: 样式信息
        """
        debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
        
        # 获取图片信息
        if not hasattr(token, 'attrs') or not token.attrs:
            if debug:
                print("警告: 图片标记没有属性")
            return
        
        # 获取图片URL和标题
        src = token.attrs.get('src', '')
        
        # 获取alt文本
        alt = ""
        if hasattr(token, 'content'):
            alt = token.content
        
        if debug:
            print(f"处理段落内图片: src={src}, alt={alt}")
        
        # 解析尺寸信息（如果有）
        width, height = self._parse_size(alt)
        if width and height and debug:
            print(f"图片尺寸: {width}x{height}")
        
        # 添加图片
        try:
            # 获取图片数据
            image_data = self._get_image_data(src)
            if not image_data:
                if debug:
                    print(f"无法获取图片数据: {src}")
                return
            
            # 添加图片到段落
            run = paragraph.add_run()
            if width and height:
                # 使用指定尺寸
                run.add_picture(image_data, width=Pt(width), height=Pt(height))
            else:
                # 使用默认尺寸（较小，适合内联）
                run.add_picture(image_data, width=Pt(100))
            
            if debug:
                print(f"段落内图片添加成功: {src}")
                
        except Exception as e:
            if debug:
                print(f"添加段落内图片失败: {str(e)}")
    
    def _get_image_data(self, src: str) -> Optional[BytesIO]:
        """获取图片数据
        
        Args:
            src: 图片路径或URL
            
        Returns:
            BytesIO: 图片数据流
        """
        # 检查缓存
        if src in self._image_cache:
            return BytesIO(self._image_cache[src])
        
        try:
            # 处理在线图片
            if src.startswith(('http://', 'https://')):
                response = requests.get(src, timeout=10)
                if response.status_code == 200:
                    image_data = response.content
                    # 缓存图片数据
                    self._image_cache[src] = image_data
                    return BytesIO(image_data)
            # 处理本地图片
            else:
                # 尝试从当前目录加载
                if os.path.exists(src):
                    with open(src, 'rb') as f:
                        image_data = f.read()
                        # 缓存图片数据
                        self._image_cache[src] = image_data
                        return BytesIO(image_data)
                
                # 尝试从测试目录加载
                test_path = os.path.join('tests', 'samples', 'basic', src)
                if os.path.exists(test_path):
                    with open(test_path, 'rb') as f:
                        image_data = f.read()
                        # 缓存图片数据
                        self._image_cache[src] = image_data
                        return BytesIO(image_data)
        except Exception as e:
            debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
            if debug:
                print(f"获取图片数据失败: {str(e)}")
        
        return None
    
    def _parse_size(self, alt: str) -> Tuple[Optional[int], Optional[int]]:
        """从alt文本中解析图片尺寸
        
        格式: ![alt|widthxheight](src)
        
        Args:
            alt: 图片alt文本
            
        Returns:
            Tuple[Optional[int], Optional[int]]: (宽度, 高度)
        """
        debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
        
        if not alt:
            return None, None
        
        # 分割alt文本和尺寸信息
        if '|' in alt:
            parts = alt.split('|')
            if len(parts) != 2:
                return None, None
            
            # 解析尺寸信息
            size_part = parts[1].strip()
            match = re.match(r'(\d+)x(\d+)', size_part)
            if match:
                width = int(match.group(1))
                height = int(match.group(2))
                if debug:
                    print(f"解析到图片尺寸: {width}x{height}")
                return width, height
        
        return None, None 