"""
表格转换器模块
"""
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .base import ElementConverter


class TableConverter(ElementConverter):
    """表格转换器，处理Markdown表格到DOCX表格的转换"""

    def __init__(self, base_converter=None):
        """初始化表格转换器
        
        Args:
            base_converter: 基础转换器实例，用于处理表格内的内联元素
        """
        super().__init__()
        self.base_converter = base_converter
        self.debug = False
        self.current_style = {}  # 当前样式
        if base_converter:
            self.debug = base_converter.debug

    def convert(self, token, tokens=None):
        """转换表格token为DOCX表格
        
        Args:
            token: 表格token
            tokens: 所有tokens列表，用于上下文处理
            
        Returns:
            docx.table: 创建的表格对象
        """
        if not self.document:
            raise ValueError("Document not set for TableConverter")
        
        if self.debug:
            print(f"处理表格: {token}")
            if tokens:
                print(f"表格tokens: {tokens}")
        
        # 解析表格结构
        rows = self._parse_table_structure(token, tokens)
        if not rows:
            return None
        
        # 获取列数
        cols = len(rows[0]) if rows else 0
        
        # 创建表格
        table = self.document.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        
        # 填充表格内容
        self._fill_table_content(table, rows)
        
        # 设置表格对齐方式
        self._set_table_alignment(table, token)
        
        return table
    
    def _parse_table_structure(self, token, tokens=None):
        """解析表格结构
        
        Args:
            token: 表格token
            tokens: 所有tokens列表，用于上下文处理
            
        Returns:
            list: 包含表格所有行和单元格内容的列表
        """
        rows = []
        
        # 如果没有提供tokens，尝试从token中获取children
        if not tokens and hasattr(token, 'children'):
            tokens = [token] + list(token.children)
        
        # 如果仍然没有tokens，返回空列表
        if not tokens:
            return []
        
        # 调试输出
        if self.debug:
            print(f"解析表格结构，tokens长度: {len(tokens)}")
            for t in tokens:
                print(f"  Token: {t.type}")
        
        # 查找表格行
        tr_open_indices = []
        tr_close_indices = []
        
        for i, t in enumerate(tokens):
            if hasattr(t, 'type'):
                if t.type == 'tr_open':
                    tr_open_indices.append(i)
                elif t.type == 'tr_close':
                    tr_close_indices.append(i)
        
        # 确保找到了相同数量的开始和结束标记
        if len(tr_open_indices) != len(tr_close_indices):
            if self.debug:
                print(f"警告: 表格行的开始和结束标记数量不匹配: {len(tr_open_indices)} vs {len(tr_close_indices)}")
            # 尝试修复
            if len(tr_open_indices) > len(tr_close_indices):
                tr_close_indices.append(len(tokens) - 1)
            else:
                tr_open_indices.append(0)
        
        # 处理每一行
        for i in range(len(tr_open_indices)):
            start_idx = tr_open_indices[i]
            end_idx = tr_close_indices[i] if i < len(tr_close_indices) else len(tokens) - 1
            
            # 提取行内容
            row_tokens = tokens[start_idx:end_idx+1]
            
            # 查找单元格
            th_open_indices = []
            th_close_indices = []
            td_open_indices = []
            td_close_indices = []
            
            for j, rt in enumerate(row_tokens):
                if hasattr(rt, 'type'):
                    if rt.type == 'th_open':
                        th_open_indices.append(j)
                    elif rt.type == 'th_close':
                        th_close_indices.append(j)
                    elif rt.type == 'td_open':
                        td_open_indices.append(j)
                    elif rt.type == 'td_close':
                        td_close_indices.append(j)
            
            # 处理表头单元格
            cells = []
            for j in range(len(th_open_indices)):
                start_j = th_open_indices[j]
                end_j = th_close_indices[j] if j < len(th_close_indices) else len(row_tokens) - 1
                
                # 提取单元格内容
                cell_tokens = row_tokens[start_j+1:end_j]
                
                # 获取对齐方式
                align = None
                if hasattr(row_tokens[start_j], 'attrs'):
                    if 'style' in row_tokens[start_j].attrs:
                        style = row_tokens[start_j].attrs['style']
                        if 'text-align:left' in style:
                            align = 'left'
                        elif 'text-align:center' in style:
                            align = 'center'
                        elif 'text-align:right' in style:
                            align = 'right'
                    
                    # 从token属性中获取align信息
                    if not align and 'align' in row_tokens[start_j].attrs:
                        align = row_tokens[start_j].attrs['align']
                
                cells.append({
                    'content': cell_tokens,
                    'is_header': True,
                    'align': align
                })
            
            # 处理数据单元格
            for j in range(len(td_open_indices)):
                start_j = td_open_indices[j]
                end_j = td_close_indices[j] if j < len(td_close_indices) else len(row_tokens) - 1
                
                # 提取单元格内容
                cell_tokens = row_tokens[start_j+1:end_j]
                
                # 获取对齐方式
                align = None
                if hasattr(row_tokens[start_j], 'attrs'):
                    if 'style' in row_tokens[start_j].attrs:
                        style = row_tokens[start_j].attrs['style']
                        if 'text-align:left' in style:
                            align = 'left'
                        elif 'text-align:center' in style:
                            align = 'center'
                        elif 'text-align:right' in style:
                            align = 'right'
                    
                    # 从token属性中获取align信息
                    if not align and 'align' in row_tokens[start_j].attrs:
                        align = row_tokens[start_j].attrs['align']
                
                cells.append({
                    'content': cell_tokens,
                    'is_header': False,
                    'align': align
                })
            
            if cells:
                rows.append(cells)
        
        # 如果没有找到行，尝试直接从token的children中提取
        if not rows and hasattr(token, 'children'):
            for row_token in token.children:
                if hasattr(row_token, 'type') and row_token.type == 'tr':
                    row = []
                    for cell_token in row_token.children:
                        if hasattr(cell_token, 'type') and cell_token.type in ('th', 'td'):
                            # 获取单元格内容
                            cell_content = cell_token.children if hasattr(cell_token, 'children') else []
                            row.append({
                                'content': cell_content,
                                'is_header': cell_token.type == 'th',
                                'align': self._get_cell_alignment(cell_token)
                            })
                    if row:
                        rows.append(row)
        
        # 调试输出
        if self.debug:
            print(f"解析到 {len(rows)} 行表格")
            for i, row in enumerate(rows):
                print(f"  行 {i+1}: {len(row)} 个单元格")
        
        return rows
    
    def _get_cell_alignment(self, cell_token):
        """获取单元格对齐方式
        
        Args:
            cell_token: 单元格token
            
        Returns:
            str: 对齐方式 ('left', 'center', 'right' 或 None)
        """
        if hasattr(cell_token, 'attrs') and 'style' in cell_token.attrs:
            style = cell_token.attrs['style']
            if 'text-align:left' in style:
                return 'left'
            elif 'text-align:center' in style:
                return 'center'
            elif 'text-align:right' in style:
                return 'right'
        
        # 从token属性中获取align信息
        if hasattr(cell_token, 'attrs') and 'align' in cell_token.attrs:
            return cell_token.attrs['align']
            
        return None
    
    def _fill_table_content(self, table, rows):
        """填充表格内容
        
        Args:
            table: docx表格对象
            rows: 解析后的表格行数据
        """
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                
                # 设置单元格垂直对齐方式
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 处理单元格内容
                if cell_data['content']:
                    # 如果有基础转换器，使用它来处理内联内容
                    if self.base_converter:
                        p = cell.paragraphs[0]
                        for content_token in cell_data['content']:
                            if hasattr(content_token, 'type') and content_token.type == 'inline':
                                # 使用基础转换器处理内联内容
                                if hasattr(self.base_converter, '_process_inline'):
                                    self.base_converter._process_inline(content_token, p)
                                # 直接处理内联内容（如果基础转换器没有_process_inline方法）
                                elif hasattr(content_token, 'children'):
                                    for child in content_token.children:
                                        if hasattr(child, 'type'):
                                            if child.type == 'text':
                                                run = p.add_run(child.content)
                                                # 应用当前样式
                                                if 'bold' in self.current_style:
                                                    run.bold = self.current_style['bold']
                                                if 'italic' in self.current_style:
                                                    run.italic = self.current_style['italic']
                                                if 'strike' in self.current_style:
                                                    run.font.strike = self.current_style['strike']
                                            elif child.type == 'strong_open':
                                                # 开始加粗
                                                self.current_style = {'bold': True}
                                            elif child.type == 'strong_close':
                                                # 结束加粗
                                                self.current_style = {}
                                            elif child.type == 'em_open':
                                                # 开始斜体
                                                self.current_style = {'italic': True}
                                            elif child.type == 'em_close':
                                                # 结束斜体
                                                self.current_style = {}
                                            elif child.type == 's_open':
                                                # 开始删除线
                                                self.current_style = {'strike': True}
                                            elif child.type == 's_close':
                                                # 结束删除线
                                                self.current_style = {}
                            elif hasattr(content_token, 'type') and content_token.type == 'text':
                                p.add_run(content_token.content)
                            else:
                                # 处理其他类型的内容
                                try:
                                    # 尝试使用基础转换器处理
                                    if hasattr(self.base_converter, '_process_token'):
                                        self.base_converter._process_token(content_token, cell)
                                    else:
                                        # 简单文本处理
                                        if hasattr(content_token, 'content'):
                                            p.add_run(content_token.content)
                                except Exception as e:
                                    if self.debug:
                                        print(f"处理单元格内容时出错: {e}")
                    else:
                        # 简单文本处理
                        text = self._get_text_from_tokens(cell_data['content'])
                        cell.text = text
                
                # 设置单元格水平对齐方式
                self._set_cell_alignment(cell, cell_data['align'])
                
                # 设置表头样式
                if cell_data['is_header']:
                    self._set_header_style(cell)
    
    def _get_text_from_tokens(self, tokens):
        """从tokens中提取文本内容
        
        Args:
            tokens: token列表
            
        Returns:
            str: 提取的文本内容
        """
        text = ""
        for token in tokens:
            if hasattr(token, 'type') and token.type == 'text':
                text += token.content
            elif hasattr(token, 'content'):
                text += token.content
            elif hasattr(token, 'children'):
                text += self._get_text_from_tokens(token.children)
        return text
    
    def _set_cell_alignment(self, cell, align):
        """设置单元格水平对齐方式
        
        Args:
            cell: docx单元格对象
            align: 对齐方式 ('left', 'center', 'right' 或 None)
        """
        if not align:
            return
            
        for paragraph in cell.paragraphs:
            if align == 'left':
                paragraph.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT
            elif align == 'center':
                paragraph.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
            elif align == 'right':
                paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    def _set_header_style(self, cell):
        """设置表头单元格样式
        
        Args:
            cell: docx表头单元格对象
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    def _set_table_alignment(self, table, token):
        """设置表格整体对齐方式
        
        Args:
            table: docx表格对象
            token: 表格token
        """
        # 默认表格宽度为页面宽度的90%
        table.autofit = False
        table.width = Inches(6)
        
        # 如果需要设置表格对齐方式，可以在这里添加代码
        # 例如居中对齐表格
        self._set_table_center_alignment(table)
    
    def _set_table_center_alignment(self, table):
        """设置表格居中对齐
        
        Args:
            table: docx表格对象
        """
        # 使用XML方式设置表格居中
        tbl = table._tbl
        tblPr = tbl.xpath('w:tblPr')[0]
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc) 