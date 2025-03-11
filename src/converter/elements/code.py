from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from .base import ElementConverter


class CodeConverter(ElementConverter):
    """代码块转换器"""

    def __init__(self, base_converter=None):
        super().__init__(base_converter)
        self.document = None
        self._last_was_code = False

    def set_document(self, document):
        if document is None:
            raise ValueError("Document cannot be None")
            
        self.document = document
        # 创建代码样式
        if 'Code' not in self.document.styles:
            style = self.document.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Consolas'  # 使用等宽字体
            font.size = Pt(10)
            # 设置段落格式
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(10)
            style.paragraph_format.left_indent = Pt(32)  # 约0.5英寸
            style.paragraph_format.right_indent = Pt(32)  # 约0.5英寸

    def convert(self, token):
        """转换代码块

        Args:
            token: 代码块标记
        """
        if not self.document:
            raise ValueError("Document not set")

        # 如果上一个是代码块，添加空行
        if self._last_was_code:
            self.document.add_paragraph()

        # 创建新段落
        paragraph = self.document.add_paragraph()
        paragraph.style = 'Code'

        # 获取代码内容
        code = token.content if hasattr(token, 'content') else ''
        
        # 处理空的代码块
        if not code:
            paragraph.add_run("")
            return

        # 分割并处理每一行，去掉末尾的空行
        lines = code.rstrip('\n').splitlines()
        
        # 添加代码内容
        for i, line in enumerate(lines):
            if i > 0:  # 不是第一行，添加换行符
                paragraph.add_run('\n')
            run = paragraph.add_run(line)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(51, 51, 51)  # 深灰色

        # 更新状态
        self._last_was_code = True 