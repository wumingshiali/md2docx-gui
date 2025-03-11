"""
文本转换器模块，处理段落和内联文本的转换
"""
from typing import Any, Dict, List, Tuple
from docx.shared import Pt
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from .base import ElementConverter


class TextConverter(ElementConverter):
    """处理段落和内联文本的转换器"""
    
    def __init__(self, base_converter=None):
        """初始化文本转换器
        
        Args:
            base_converter: 基础转换器实例，用于访问其他转换器
        """
        super().__init__(base_converter)
        self.document = None
        
    def convert(self, tokens: Tuple[Any, Any]) -> None:
        """转换段落元素
        
        Args:
            tokens: (开始标记, 内容标记) 的元组
        """
        if not self.document:
            raise ValueError("Document not set")
            
        paragraph_token, content_token = tokens
        
        # 创建新段落
        paragraph = self.document.add_paragraph()
        
        # 处理空段落
        if not content_token or not hasattr(content_token, 'children') or not content_token.children:
            paragraph.add_run("")
            return
        
        # 调试信息：打印段落内容
        debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
        if debug:
            print(f"处理段落: {content_token.content}")
        
        # 检查是否有链接转换器
        link_converter = None
        if self.base_converter and 'link' in self.base_converter.converters:
            link_converter = self.base_converter.converters.get('link')
            if debug:
                print(f"找到链接转换器: {link_converter}")
        else:
            if debug:
                print("未找到链接转换器")
        
        # 检查是否有图片转换器
        image_converter = None
        if self.base_converter and 'image' in self.base_converter.converters:
            image_converter = self.base_converter.converters.get('image')
            if debug:
                print(f"找到图片转换器: {image_converter}")
        else:
            if debug:
                print("未找到图片转换器")
        
        # 处理段落内的文本和样式
        current_text = ""
        current_style = {"bold": False, "italic": False, "strike": False}
        
        i = 0
        children = content_token.children
        while i < len(children):
            child = children[i]
            
            # 处理标记
            if debug:
                print(f"处理标记: type={child.type}, content={child.content if hasattr(child, 'content') else ''}")
            
            if child.type == 'text':
                # 检查是否是带样式的链接的一部分
                is_styled_link = False
                
                # 检查是否是粗体链接的开始
                if child.content.endswith('**') and i + 1 < len(children) and children[i + 1].type == 'link_open':
                    if debug:
                        print("检测到粗体链接开始")
                    # 移除末尾的 **
                    text = child.content[:-2]
                    current_text += text
                    
                    # 添加当前文本
                    if current_text:
                        self._add_text_with_style(paragraph, current_text, current_style)
                        current_text = ""
                    
                    # 设置粗体样式
                    current_style["bold"] = True
                    
                    # 处理链接
                    link_token = children[i + 1]
                    link_content = None
                    j = i + 2
                    while j < len(children) and children[j].type != 'link_close':
                        if children[j].type == 'text':
                            link_content = children[j]
                        j += 1
                    
                    # 如果找到链接内容
                    if link_content and link_converter:
                        if debug:
                            print(f"处理粗体链接: {link_content.content}")
                        # 传递链接文本
                        link_text = link_content.content if hasattr(link_content, 'content') else None
                        link_converter.convert_in_paragraph(paragraph, link_token, current_style.copy(), link_text)
                    else:
                        # 如果没有找到链接转换器，使用普通文本
                        if link_content:
                            self._add_text_with_style(paragraph, link_content.content, current_style)
                    
                    # 跳过已处理的标记
                    i = j + 1 if j < len(children) else i + 1
                    
                    # 检查下一个标记是否是粗体链接的结束
                    if i < len(children) and children[i].type == 'text' and children[i].content.startswith('**'):
                        if debug:
                            print("检测到粗体链接结束")
                        # 移除开头的 **
                        next_text = children[i].content[2:]
                        
                        # 重置粗体样式
                        current_style["bold"] = False
                        
                        # 如果还有其他文本，添加到当前文本
                        if next_text:
                            current_text = next_text
                        
                        i += 1
                    
                    is_styled_link = True
                
                # 检查是否是斜体链接的开始
                elif child.content.endswith('*') and not child.content.endswith('**') and i + 1 < len(children) and children[i + 1].type == 'link_open':
                    if debug:
                        print("检测到斜体链接开始")
                    # 移除末尾的 *
                    text = child.content[:-1]
                    current_text += text
                    
                    # 添加当前文本
                    if current_text:
                        self._add_text_with_style(paragraph, current_text, current_style)
                        current_text = ""
                    
                    # 设置斜体样式
                    current_style["italic"] = True
                    
                    # 处理链接
                    link_token = children[i + 1]
                    link_content = None
                    j = i + 2
                    while j < len(children) and children[j].type != 'link_close':
                        if children[j].type == 'text':
                            link_content = children[j]
                        j += 1
                    
                    # 如果找到链接内容
                    if link_content and link_converter:
                        if debug:
                            print(f"处理斜体链接: {link_content.content}")
                        # 传递链接文本
                        link_text = link_content.content if hasattr(link_content, 'content') else None
                        link_converter.convert_in_paragraph(paragraph, link_token, current_style.copy(), link_text)
                    else:
                        # 如果没有找到链接转换器，使用普通文本
                        if link_content:
                            self._add_text_with_style(paragraph, link_content.content, current_style)
                    
                    # 跳过已处理的标记
                    i = j + 1 if j < len(children) else i + 1
                    
                    # 检查下一个标记是否是斜体链接的结束
                    if i < len(children) and children[i].type == 'text' and children[i].content.startswith('*') and not children[i].content.startswith('**'):
                        if debug:
                            print("检测到斜体链接结束")
                        # 移除开头的 *
                        next_text = children[i].content[1:]
                        
                        # 重置斜体样式
                        current_style["italic"] = False
                        
                        # 如果还有其他文本，添加到当前文本
                        if next_text:
                            current_text = next_text
                        
                        i += 1
                    
                    is_styled_link = True
                
                # 检查是否是删除线链接的开始
                elif child.content.endswith('~~') and i + 1 < len(children) and children[i + 1].type == 'link_open':
                    if debug:
                        print("检测到删除线链接开始")
                    # 移除末尾的 ~~
                    text = child.content[:-2]
                    current_text += text
                    
                    # 添加当前文本
                    if current_text:
                        self._add_text_with_style(paragraph, current_text, current_style)
                        current_text = ""
                    
                    # 设置删除线样式
                    current_style["strike"] = True
                    
                    # 处理链接
                    link_token = children[i + 1]
                    link_content = None
                    j = i + 2
                    while j < len(children) and children[j].type != 'link_close':
                        if children[j].type == 'text':
                            link_content = children[j]
                        j += 1
                    
                    # 如果找到链接内容
                    if link_content and link_converter:
                        if debug:
                            print(f"处理删除线链接: {link_content.content}")
                        # 传递链接文本
                        link_text = link_content.content if hasattr(link_content, 'content') else None
                        link_converter.convert_in_paragraph(paragraph, link_token, current_style.copy(), link_text)
                    else:
                        # 如果没有找到链接转换器，使用普通文本
                        if link_content:
                            self._add_text_with_style(paragraph, link_content.content, current_style)
                    
                    # 跳过已处理的标记
                    i = j + 1 if j < len(children) else i + 1
                    
                    # 检查下一个标记是否是删除线链接的结束
                    if i < len(children) and children[i].type == 'text' and children[i].content.startswith('~~'):
                        if debug:
                            print("检测到删除线链接结束")
                        # 移除开头的 ~~
                        next_text = children[i].content[2:]
                        
                        # 重置删除线样式
                        current_style["strike"] = False
                        
                        # 如果还有其他文本，添加到当前文本
                        if next_text:
                            current_text = next_text
                        
                        i += 1
                    
                    is_styled_link = True
                
                # 如果不是带样式的链接，正常处理文本
                if not is_styled_link:
                    # 处理多行文本中的空格
                    text = child.content.replace('\n', ' ')
                    if text.endswith(' '):
                        text = text[:-1]
                    current_text += text
                    i += 1
            elif child.type == 'link_open':
                # 处理链接前的文本
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                
                # 获取链接内容
                link_content = None
                j = i + 1
                while j < len(children) and children[j].type != 'link_close':
                    if children[j].type == 'text':
                        link_content = children[j]
                    j += 1
                
                # 处理链接
                if link_content and link_converter:
                    if debug:
                        print(f"处理普通链接: {link_content.content}")
                    # 传递链接文本
                    link_text = link_content.content if hasattr(link_content, 'content') else None
                    link_converter.convert_in_paragraph(paragraph, child, current_style.copy(), link_text)
                else:
                    # 如果没有找到链接转换器，使用普通文本
                    if link_content:
                        self._add_text_with_style(paragraph, link_content.content, current_style)
                
                # 跳过已处理的标记
                i = j + 1 if j < len(children) else i + 1
            elif child.type == 'link_close':
                i += 1
            elif child.type == 'image':
                # 处理图片前的文本
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                
                # 处理图片
                if image_converter:
                    if debug:
                        print(f"处理段落内图片")
                    image_converter.convert_in_paragraph(paragraph, child, current_style.copy())
                
                i += 1
            elif child.type == 'strong_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["bold"] = True
                i += 1
            elif child.type == 'strong_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["bold"] = False
                i += 1
            elif child.type == 'em_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["italic"] = True
                i += 1
            elif child.type == 'em_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["italic"] = False
                i += 1
            elif child.type == 's_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["strike"] = True
                i += 1
            elif child.type == 's_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["strike"] = False
                i += 1
            elif child.type == 'softbreak':
                current_text += " "
                i += 1
            else:
                i += 1
        
        # 添加剩余的文本
        if current_text:
            self._add_text_with_style(paragraph, current_text, current_style)
    
    def _add_text_with_style(self, paragraph: Paragraph, text: str, style: Dict[str, bool]) -> None:
        """添加带样式的文本
        
        Args:
            paragraph: 段落对象
            text: 要添加的文本
            style: 样式配置
        """
        run = paragraph.add_run(text)
        run.bold = style["bold"]
        run.italic = style["italic"]
        run.font.strike = style["strike"]
    
    def _get_text_between_tokens(self, tokens: List[Any], start_token: Any) -> str:
        """获取开始和结束标记之间的文本
        
        Args:
            tokens: 标记列表
            start_token: 开始标记
            
        Returns:
            str: 标记之间的文本
        """
        start_idx = tokens.index(start_token)
        text_token = tokens[start_idx + 1]
        return text_token.content if text_token.type == 'text' else '' 