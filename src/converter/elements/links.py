"""
链接转换器模块
"""
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .base import ElementConverter


class LinkConverter(ElementConverter):
    """链接转换器，处理各种类型的链接"""

    def __init__(self, base_converter=None):
        super().__init__(base_converter)
        self.document = None

    def set_document(self, document):
        if document is None:
            raise ValueError("Document cannot be None")
        self.document = document
        # 创建链接样式
        self._ensure_hyperlink_style()

    def _ensure_hyperlink_style(self):
        """确保Hyperlink样式存在"""
        if 'Hyperlink' not in self.document.styles:
            style = self.document.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
            font = style.font
            font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            font.underline = True

    def convert(self, token_pair):
        """转换链接

        Args:
            token_pair: (token, content_token) 元组，包含链接标记和内容
        """
        if not self.document:
            raise ValueError("Document not set")

        token, content_token = token_pair
        if not hasattr(token, 'attrs'):
            return

        # 获取链接URL和标题
        url = token.attrs.get('href', '') if hasattr(token, 'attrs') else ''
        
        # 设置链接文本
        text = content_token.content if hasattr(content_token, 'content') else url
        
        # 如果文本为空，使用URL作为文本
        if not text and url:
            text = url
            
        # 确保有文本内容
        if not text:
            text = "(空链接)"
        
        # 获取当前段落或创建新段落
        if self.document.paragraphs:
            paragraph = self.document.paragraphs[-1]
        else:
            paragraph = self.document.add_paragraph()
        
        # 创建超链接
        self._add_hyperlink(paragraph, text, url)
    
    def convert_in_paragraph(self, paragraph, token, style=None, link_text=None):
        """在段落中转换链接
        
        Args:
            paragraph: 段落对象
            token: 链接标记
            style: 样式信息
            link_text: 链接文本，如果提供则使用此文本
        """
        debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
        if debug:
            print(f"转换链接: token={token.type}, content={token.content if hasattr(token, 'content') else ''}")
            print(f"链接样式: {style}")
        
        if not hasattr(token, 'attrs') or not token.attrs:
            if debug:
                print("警告: 链接标记没有属性")
            return
            
        # 获取链接URL
        url = token.attrs.get('href', '')
        if debug:
            print(f"链接URL: {url}")
        
        # 获取链接文本
        text = link_text or ""
        
        # 如果没有提供链接文本，尝试从token中获取
        if not text and hasattr(token, 'children') and token.children:
            for child in token.children:
                if child.type == 'text':
                    text += child.content
        
        if debug:
            print(f"链接文本: {text}")
        
        # 如果没有文本，使用URL作为文本
        if not text:
            text = url
            if debug:
                print(f"使用URL作为链接文本: {text}")
            
        # 添加带样式的超链接
        self._add_hyperlink_with_style(paragraph, text, url, style or {})
    
    def _add_hyperlink(self, paragraph, text, url):
        """添加超链接到段落
        
        Args:
            paragraph: 段落对象
            text: 链接文本
            url: 链接地址
        """
        # 创建超链接
        run = paragraph.add_run(text)
        
        # 确保Hyperlink样式存在
        self._ensure_hyperlink_style()
        
        # 应用超链接样式
        run.style = 'Hyperlink'
        
        # 如果URL为空，不创建实际的超链接
        if not url:
            return
        
        # 创建关系ID
        part = self.document.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # 获取XML元素
        r_element = run._element
        
        # 创建超链接XML元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 获取父元素
        parent = r_element.getparent()
        
        # 获取运行元素的索引
        index = parent.index(r_element)
        
        # 从父元素中移除运行元素
        parent.remove(r_element)
        
        # 将运行元素添加到超链接
        hyperlink.append(r_element)
        
        # 将超链接插入到原来运行元素的位置
        parent.insert(index, hyperlink)
    
    def _add_hyperlink_with_style(self, paragraph, text, url, style):
        """添加带样式的超链接到段落
        
        Args:
            paragraph: 段落对象
            text: 链接文本
            url: 链接地址
            style: 样式信息，包含bold、italic、strike
        """
        debug = self.base_converter.debug if hasattr(self.base_converter, 'debug') else False
        # 调试信息
        if debug:
            print(f"添加带样式的超链接: text='{text}', url='{url}', style={style}")
        
        # 创建超链接
        run = paragraph.add_run(text)
        if debug:
            print(f"创建的run文本: '{run.text}'")
        
        # 应用样式
        if style.get("bold"):
            run.bold = True
            if debug:
                print("应用粗体样式")
        if style.get("italic"):
            run.italic = True
            if debug:
                print("应用斜体样式")
        if style.get("strike"):
            run.font.strike = True
            if debug:
                print("应用删除线样式")
        
        # 确保Hyperlink样式存在
        self._ensure_hyperlink_style()
        
        # 应用超链接样式
        run.style = 'Hyperlink'
        
        # 如果URL为空，不创建实际的超链接
        if not url:
            if debug:
                print("URL为空，不创建实际的超链接")
            return
        
        # 创建关系ID
        part = self.document.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # 获取XML元素
        r_element = run._element
        
        # 创建超链接XML元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 获取父元素
        parent = r_element.getparent()
        
        # 获取运行元素的索引
        index = parent.index(r_element)
        
        # 从父元素中移除运行元素
        parent.remove(r_element)
        
        # 将运行元素添加到超链接
        hyperlink.append(r_element)
        
        # 将超链接插入到原来运行元素的位置
        parent.insert(index, hyperlink)
        
        if debug:
            print("超链接创建成功") 