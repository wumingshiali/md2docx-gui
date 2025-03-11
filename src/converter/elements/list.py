"""
列表转换器模块，处理有序列表和无序列表的转换
"""
from typing import Any, Dict, Tuple, List, Optional
from docx.shared import Pt, Inches
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import ns
from .base import ElementConverter


class ListConverter(ElementConverter):
    """处理列表的转换器"""
    
    def __init__(self, base_converter=None):
        super().__init__(base_converter)
        # 跟踪当前列表状态：(层级, 是否有序, 编号ID)
        self._current_lists: List[Tuple[int, bool, Optional[int]]] = []
        # 缓存已创建的编号定义：(层级, 是否有序) -> 编号ID
        self._numbering_cache: Dict[Tuple[int, bool], int] = {}
        # 跟踪每个层级的当前编号
        self._current_numbers: Dict[int, int] = {}
        # 上一个处理的标记类型
        self._last_token_type: Optional[str] = None
    
    def convert(self, tokens: Tuple[Any, Any]) -> Paragraph:
        """转换列表元素
        
        Args:
            tokens: (开始标记, 内容标记) 的元组
            
        Returns:
            Paragraph: 创建的段落对象
        """
        if not self.document:
            raise ValueError("Document not set")
            
        list_token, content_token = tokens
        
        # 获取列表层级和类型
        level, is_ordered = self._get_list_info(list_token)
        
        # 检查是否需要重新开始编号
        need_new_numbering = self._check_if_need_new_numbering(level, is_ordered)
        
        # 如果需要重新编号，重置该层级的编号
        if need_new_numbering:
            self._current_numbers[level] = 1
        elif is_ordered:
            # 如果不需要重新编号，且是有序列表，增加该层级的编号
            if level not in self._current_numbers:
                self._current_numbers[level] = 1
            else:
                self._current_numbers[level] += 1
        
        # 更新上一个标记类型
        self._last_token_type = list_token.type
        
        # 创建或获取列表样式
        style_name = self._get_style_name(level, is_ordered)
        numbering_id = self._ensure_list_style(style_name, level, is_ordered, need_new_numbering)
        
        # 更新列表状态
        self._update_list_state(level, is_ordered, numbering_id)
        
        # 创建新段落
        paragraph = self.document.add_paragraph()
        paragraph.style = self.document.styles[style_name]
        
        # 处理列表项内的文本和样式
        current_text = ""
        current_style = {"bold": False, "italic": False, "strike": False}
        
        # 处理空列表项
        if not content_token or not hasattr(content_token, 'children') or not content_token.children:
            # 检查是否为任务列表项（通过内容字符串判断）
            if hasattr(content_token, 'content') and isinstance(content_token.content, str):
                content = content_token.content.strip()
                if content.startswith('[ ] ') or content.startswith('[x] '):
                    # 这是一个任务列表项，但我们在这里不处理它的内容
                    # 只返回一个空段落，让任务列表转换器处理内容
                    return paragraph
                else:
                    # 普通列表项，添加内容
                    paragraph.add_run(content)
                    return paragraph
            else:
                # 真正的空列表项
                paragraph.add_run("")
                return paragraph
        
        # 检查是否为任务列表项（通过内容字符串判断）
        if hasattr(content_token, 'content') and isinstance(content_token.content, str):
            content = content_token.content.strip()
            if content.startswith('[ ] ') or content.startswith('[x] '):
                # 这是一个任务列表项，但我们在这里不处理它的内容
                # 只返回一个空段落，让任务列表转换器处理内容
                return paragraph
        
        # 处理列表项内的子元素
        for child in content_token.children:
            if child.type == 'text':
                # 处理多行文本中的空格
                text = child.content.replace('\n', ' ')
                if text.endswith(' '):
                    text = text[:-1]
                current_text += text
            elif child.type == 'strong_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["bold"] = True
            elif child.type == 'strong_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["bold"] = False
            elif child.type == 'em_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["italic"] = True
            elif child.type == 'em_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["italic"] = False
            elif child.type == 's_open':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["strike"] = True
            elif child.type == 's_close':
                if current_text:
                    self._add_text_with_style(paragraph, current_text, current_style)
                    current_text = ""
                current_style["strike"] = False
            elif child.type == 'softbreak':
                current_text += " "
        
        # 添加剩余的文本
        if current_text:
            self._add_text_with_style(paragraph, current_text, current_style)
            
        return paragraph
    
    def _check_if_need_new_numbering(self, level: int, is_ordered: bool) -> bool:
        """检查是否需要重新开始编号
        
        Args:
            level: 列表层级
            is_ordered: 是否为有序列表
            
        Returns:
            bool: 是否需要重新开始编号
        """
        # 如果是无序列表，不需要重新编号
        if not is_ordered:
            return False
            
        # 如果列表状态为空，需要新编号
        if not self._current_lists:
            return True
            
        # 如果上一个标记是标题，需要重新编号
        if self._last_token_type and self._last_token_type.endswith('heading_close'):
            return True
            
        # 检查是否有相同层级的列表
        for l, o, _ in reversed(self._current_lists):
            if l == level:
                # 如果找到相同层级，但类型不同，需要新编号
                if o != is_ordered:
                    return True
                # 如果类型相同，不需要新编号
                return False
            # 如果遇到更深层级，继续查找
            if l > level:
                continue
            # 如果遇到更浅层级，需要新编号
            if l < level:
                return True
            
        return True
    
    def _update_list_state(self, level: int, is_ordered: bool, numbering_id: Optional[int]) -> None:
        """更新列表状态
        
        Args:
            level: 列表层级
            is_ordered: 是否为有序列表
            numbering_id: 编号定义ID
        """
        # 移除所有更深层级的列表状态
        while self._current_lists and self._current_lists[-1][0] > level:
            self._current_lists.pop()
            
        # 如果当前层级已存在，更新它
        if self._current_lists and self._current_lists[-1][0] == level:
            self._current_lists[-1] = (level, is_ordered, numbering_id)
        else:
            # 否则添加新的列表状态
            self._current_lists.append((level, is_ordered, numbering_id))
    
    def _add_text_with_style(self, paragraph: Paragraph, text: str, style: Dict[str, bool]) -> None:
        """添加带样式的文本
        
        Args:
            paragraph: 段落对象
            text: 要添加的文本
            style: 样式配置
        """
        run = paragraph.add_run(text)
        run.bold = style["bold"]
        run.italic = style["italic"]
        if style["strike"]:
            rPr = run._element.get_or_add_rPr()
            strike = OxmlElement('w:strike')
            strike.set(qn('w:val'), 'true')
            rPr.append(strike)
    
    def _get_list_info(self, token: Any) -> Tuple[int, bool]:
        """获取列表的层级和类型
        
        Args:
            token: 列表标记
            
        Returns:
            Tuple[int, bool]: (层级, 是否为有序列表)
        """
        level = 1
        is_ordered = False
        
        if hasattr(token, 'type'):
            # 判断列表类型
            is_ordered = token.type == 'ordered_list_open'
            
            # 计算缩进层级
            if hasattr(token, 'content'):
                # 通过内容的缩进计算层级
                indent = len(token.content)
                level = (indent // 2) + 1
        
        return level, is_ordered
    
    def _get_style_name(self, level: int, is_ordered: bool) -> str:
        """获取列表样式名称
        
        Args:
            level: 列表层级
            is_ordered: 是否为有序列表
            
        Returns:
            str: 样式名称
        """
        base_name = "List Number" if is_ordered else "List Bullet"
        return f"{base_name} {level}" if level > 1 else base_name
    
    def _ensure_list_style(self, style_name: str, level: int, is_ordered: bool, need_new_numbering: bool = False) -> Optional[int]:
        """确保列表样式存在
        
        Args:
            style_name: 样式名称
            level: 列表层级
            is_ordered: 是否为有序列表
            need_new_numbering: 是否需要新的编号定义
            
        Returns:
            Optional[int]: 编号定义ID
        """
        numbering_id = None
        
        # 检查样式是否已存在
        style_exists = style_name in self.document.styles
        if not style_exists:
            style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            # 设置基本样式
            style.font.size = Pt(12)
            # 根据层级设置左缩进
            style.paragraph_format.left_indent = Inches(0.5 * (level - 1))  # 修改缩进计算方式
            style.paragraph_format.first_line_indent = Inches(-0.25)  # 悬挂缩进
            # 设置段落间距
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            # 设置对齐方式
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            style = self.document.styles[style_name]
            
        # 处理编号
        # 获取缓存的编号定义
        cache_key = (level, is_ordered)
        if not need_new_numbering and cache_key in self._numbering_cache:
            numbering_id = self._numbering_cache[cache_key]
        else:
            # 创建新的编号定义
            numbering = self.document.part.numbering_part.numbering_definitions._numbering
            
            # 创建抽象编号定义
            abstract_num = OxmlElement('w:abstractNum')
            abstract_num_id = str(len(self._numbering_cache))  # 使用缓存长度作为ID
            abstract_num.set(qn('w:abstractNumId'), abstract_num_id)
            
            # 为每个层级创建编号格式
            for i in range(level):
                lvl = OxmlElement('w:lvl')
                lvl.set(qn('w:ilvl'), str(i))
                
                # 设置编号格式
                num_fmt = OxmlElement('w:numFmt')
                if is_ordered:
                    num_fmt.set(qn('w:val'), 'decimal')
                else:
                    num_fmt.set(qn('w:val'), 'bullet')
                lvl.append(num_fmt)
                
                # 设置编号文本
                lvl_text = OxmlElement('w:lvlText')
                if is_ordered:
                    # 使用当前层级的编号
                    lvl_text.set(qn('w:val'), f'%{i+1}.')
                else:
                    # 根据层级设置不同的项目符号
                    bullets = ['•', '○', '▪']
                    bullet = bullets[i % len(bullets)]
                    lvl_text.set(qn('w:val'), bullet)
                lvl.append(lvl_text)
                
                # 设置对齐方式
                lvl_jc = OxmlElement('w:lvlJc')
                lvl_jc.set(qn('w:val'), 'left')
                lvl.append(lvl_jc)
                
                # 设置缩进
                pPr = OxmlElement('w:pPr')
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), str(720 * (i + 1)))  # 720 twips = 0.5 inch
                ind.set(qn('w:hanging'), '360')  # 360 twips = 0.25 inch
                pPr.append(ind)
                lvl.append(pPr)
                
                # 设置后缀
                suff = OxmlElement('w:suff')
                suff.set(qn('w:val'), 'space')
                lvl.append(suff)
                
                # 添加级别定义
                abstract_num.append(lvl)
            
            # 添加抽象编号定义
            numbering.append(abstract_num)
            
            # 创建编号实例
            num = OxmlElement('w:num')
            num_id = str(len(self._numbering_cache) + 1)
            num.set(qn('w:numId'), num_id)
            abstract_num_id_element = OxmlElement('w:abstractNumId')
            abstract_num_id_element.set(qn('w:val'), abstract_num_id)
            num.append(abstract_num_id_element)
            
            # 添加编号实例
            numbering.append(num)
            numbering_id = int(num_id)
            
            # 缓存编号定义
            self._numbering_cache[cache_key] = numbering_id

        # 应用编号定义到样式
        if numbering_id is not None:
            style._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = numbering_id
            style._element.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level - 1

        return numbering_id 