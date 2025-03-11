"""
HTML转换器模块，处理Markdown中的HTML标签
"""
from typing import Any, Dict, List, Optional, Union
import os
import tempfile
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

try:
    import html2docx
    HTML2DOCX_AVAILABLE = True
except ImportError:
    HTML2DOCX_AVAILABLE = False

from .base import ElementConverter


class HtmlConverter(ElementConverter):
    """HTML转换器，处理Markdown中的HTML标签"""
    
    def __init__(self, base_converter=None):
        """初始化HTML转换器
        
        Args:
            base_converter: 基础转换器实例
        """
        super().__init__(base_converter)
        self.debug = False
        if base_converter:
            self.debug = base_converter.debug
    
    def convert(self, token: Any) -> Any:
        """转换HTML标签为DOCX元素
        
        Args:
            token: HTML标签token
            
        Returns:
            Any: 转换后的DOCX元素
        """
        if not self.document:
            raise ValueError("Document not set for HtmlConverter")
        
        if self.debug:
            print(f"处理HTML标签: {token.content if hasattr(token, 'content') else ''}")
        
        # 获取HTML内容
        html_content = ""
        if hasattr(token, 'content'):
            html_content = token.content
        elif hasattr(token, 'children'):
            for child in token.children:
                if hasattr(child, 'content'):
                    html_content += child.content
        
        if not html_content:
            if self.debug:
                print("HTML内容为空")
            return None
        
        # 首先尝试使用自定义解析方法
        result = self._custom_html_convert(html_content)
        if result:
            if self.debug:
                print("使用自定义HTML解析成功")
            return result
        
        # 如果自定义解析失败，尝试使用html2docx
        if HTML2DOCX_AVAILABLE:
            try:
                if self.debug:
                    print("尝试使用html2docx转换")
                
                # 创建完整的HTML文档
                full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>HTML转换</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; }}
                        table {{ border-collapse: collapse; width: 100%; }}
                        th, td {{ border: 1px solid #ddd; padding: 8px; }}
                        th {{ background-color: #f2f2f2; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                # 创建临时HTML文件
                with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
                    f.write(full_html)
                    temp_html_path = f.name
                
                if self.debug:
                    print(f"创建临时HTML文件: {temp_html_path}")
                    print(f"HTML内容: {full_html[:100]}...")
                
                # 创建临时DOCX文件路径
                temp_docx_path = temp_html_path.replace('.html', '.docx')
                
                # 使用html2docx转换
                html2docx.convert(temp_html_path, temp_docx_path)
                
                if self.debug:
                    print(f"转换完成，临时DOCX文件: {temp_docx_path}")
                    if os.path.exists(temp_docx_path):
                        print(f"临时DOCX文件大小: {os.path.getsize(temp_docx_path)} 字节")
                    else:
                        print("临时DOCX文件不存在")
                
                # 打开生成的DOCX文件
                temp_doc = Document(temp_docx_path)
                
                if self.debug:
                    print(f"临时文档包含 {len(temp_doc.paragraphs)} 个段落")
                
                # 将临时文档的内容复制到当前文档
                for paragraph in temp_doc.paragraphs:
                    if not paragraph.text.strip():
                        continue  # 跳过空段落
                        
                    p = self.document.add_paragraph()
                    for run in paragraph.runs:
                        r = p.add_run(run.text)
                        r.bold = run.bold
                        r.italic = run.italic
                        r.underline = run.underline
                        # 复制其他样式...
                
                # 复制表格
                for table in temp_doc.tables:
                    if self.debug:
                        print(f"复制表格: {len(table.rows)}行 x {len(table.columns)}列")
                    
                    new_table = self.document.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = 'Table Grid'
                    
                    # 复制单元格内容
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text
                
                # 清理临时文件
                try:
                    os.remove(temp_html_path)
                    os.remove(temp_docx_path)
                    if self.debug:
                        print("临时文件已清理")
                except Exception as e:
                    if self.debug:
                        print(f"清理临时文件失败: {e}")
                
                if self.debug:
                    print(f"HTML转换完成，添加了{len(temp_doc.paragraphs)}个段落和{len(temp_doc.tables)}个表格")
                
                # 返回最后一个添加的段落
                return self.document.paragraphs[-1] if self.document.paragraphs else None
                
            except Exception as e:
                if self.debug:
                    print(f"HTML转换失败: {e}")
                # 失败时回退到基本转换
                return self._fallback_convert(html_content)
        else:
            # html2docx不可用时回退到基本转换
            if self.debug:
                print("html2docx不可用，使用基本转换")
            return self._fallback_convert(html_content)
    
    def _custom_html_convert(self, html_content: str) -> Optional[Paragraph]:
        """自定义HTML解析和转换
        
        Args:
            html_content: HTML内容
            
        Returns:
            Optional[Paragraph]: 创建的段落，如果无法解析则返回None
        """
        try:
            if self.debug:
                print("使用自定义HTML解析")
            
            # 简单的HTML标签解析
            # 处理简单的HTML段落
            if re.match(r'^\s*<p>(.*?)</p>\s*$', html_content, re.DOTALL):
                content = re.sub(r'^\s*<p>(.*?)</p>\s*$', r'\1', html_content, flags=re.DOTALL)
                paragraph = self.document.add_paragraph()
                
                # 处理内部标签
                content = self._process_inline_tags(content, paragraph)
                
                if self.debug:
                    print(f"解析段落: {content}")
                
                return paragraph
            
            # 处理简单的div
            if re.match(r'^\s*<div[^>]*>(.*?)</div>\s*$', html_content, re.DOTALL):
                content = re.sub(r'^\s*<div[^>]*>(.*?)</div>\s*$', r'\1', html_content, flags=re.DOTALL)
                paragraph = self.document.add_paragraph()
                
                # 处理内部标签
                content = self._process_inline_tags(content, paragraph)
                
                if self.debug:
                    print(f"解析div: {content}")
                
                return paragraph
            
            # 处理简单的无序列表
            if re.match(r'^\s*<ul[^>]*>(.*?)</ul>\s*$', html_content, re.DOTALL):
                list_content = re.sub(r'^\s*<ul[^>]*>(.*?)</ul>\s*$', r'\1', html_content, flags=re.DOTALL)
                list_items = re.findall(r'<li[^>]*>(.*?)</li>', list_content, re.DOTALL)
                
                if self.debug:
                    print(f"解析无序列表: {len(list_items)}项")
                
                for item in list_items:
                    paragraph = self.document.add_paragraph(style='List Bullet')
                    self._process_inline_tags(item, paragraph)
                
                return self.document.paragraphs[-1] if self.document.paragraphs else None
            
            # 处理简单的有序列表
            if re.match(r'^\s*<ol[^>]*>(.*?)</ol>\s*$', html_content, re.DOTALL):
                list_content = re.sub(r'^\s*<ol[^>]*>(.*?)</ol>\s*$', r'\1', html_content, flags=re.DOTALL)
                list_items = re.findall(r'<li[^>]*>(.*?)</li>', list_content, re.DOTALL)
                
                if self.debug:
                    print(f"解析有序列表: {len(list_items)}项")
                
                for item in list_items:
                    paragraph = self.document.add_paragraph(style='List Number')
                    self._process_inline_tags(item, paragraph)
                
                return self.document.paragraphs[-1] if self.document.paragraphs else None
            
            # 处理简单的表格
            if re.match(r'^\s*<table[^>]*>(.*?)</table>\s*$', html_content, re.DOTALL):
                table_content = re.sub(r'^\s*<table[^>]*>(.*?)</table>\s*$', r'\1', html_content, flags=re.DOTALL)
                
                # 提取行
                rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_content, re.DOTALL)
                
                if not rows:
                    return None
                
                if self.debug:
                    print(f"解析表格: {len(rows)}行")
                
                # 计算列数
                first_row = rows[0]
                header_cells = re.findall(r'<th[^>]*>(.*?)</th>', first_row, re.DOTALL)
                data_cells = re.findall(r'<td[^>]*>(.*?)</td>', first_row, re.DOTALL)
                cols = max(len(header_cells), len(data_cells))
                
                if cols == 0:
                    return None
                
                # 创建表格
                table = self.document.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                
                # 填充表格内容
                for i, row_html in enumerate(rows):
                    # 提取单元格
                    header_cells = re.findall(r'<th[^>]*>(.*?)</th>', row_html, re.DOTALL)
                    data_cells = re.findall(r'<td[^>]*>(.*?)</td>', row_html, re.DOTALL)
                    
                    # 合并单元格列表
                    cells = header_cells + data_cells
                    
                    for j, cell_content in enumerate(cells):
                        if j < cols and i < len(table.rows):
                            # 清理单元格内容中的HTML标签
                            clean_content = re.sub(r'<[^>]*>', ' ', cell_content)
                            table.cell(i, j).text = clean_content.strip()
                
                # 添加一个空段落，以便返回
                return self.document.add_paragraph()
            
            # 无法解析，返回None
            return None
        except Exception as e:
            if self.debug:
                print(f"错误: 自定义HTML解析失败: {e}")
            return None
    
    def _process_inline_tags(self, content: str, paragraph: Paragraph) -> str:
        """处理内联HTML标签
        
        Args:
            content: HTML内容
            paragraph: 要添加内容的段落
            
        Returns:
            str: 处理后的内容
        """
        try:
            # 处理粗体
            bold_parts = re.split(r'(<strong>|</strong>)', content)
            is_bold = False
            
            # 处理斜体
            italic_parts = []
            for part in bold_parts:
                if part == '<strong>':
                    is_bold = True
                    continue
                elif part == '</strong>':
                    is_bold = False
                    continue
                
                # 分割斜体
                sub_parts = re.split(r'(<em>|</em>)', part)
                for sub_part in sub_parts:
                    italic_parts.append((sub_part, is_bold))
            
            is_italic = False
            
            # 处理下划线
            underline_parts = []
            for part, bold in italic_parts:
                if part == '<em>':
                    is_italic = True
                    continue
                elif part == '</em>':
                    is_italic = False
                    continue
                
                # 分割下划线
                sub_parts = re.split(r'(<u>|</u>)', part)
                for sub_part in sub_parts:
                    underline_parts.append((sub_part, bold, is_italic))
            
            is_underline = False
            
            # 处理删除线
            strike_parts = []
            for part, bold, italic in underline_parts:
                if part == '<u>':
                    is_underline = True
                    continue
                elif part == '</u>':
                    is_underline = False
                    continue
                
                # 分割删除线
                sub_parts = re.split(r'(<s>|</s>)', part)
                for sub_part in sub_parts:
                    strike_parts.append((sub_part, bold, italic, is_underline))
            
            is_strike = False
            
            # 添加文本
            for part, bold, italic, underline in strike_parts:
                if part == '<s>':
                    is_strike = True
                    continue
                elif part == '</s>':
                    is_strike = False
                    continue
                
                # 移除其他HTML标签
                clean_text = re.sub(r'<[^>]*>', ' ', part)
                if clean_text:
                    run = paragraph.add_run(clean_text)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    
                    # 添加删除线 - 修复方法
                    if is_strike:
                        try:
                            # 方法1：使用font属性
                            run.font.strike = True
                        except Exception as e:
                            if self.debug:
                                print(f"无法设置删除线(方法1): {e}")
                            try:
                                # 方法2：使用XML元素
                                run._element.get_or_add_rPr().set(qn('w:strike'), 'true')
                            except Exception as e:
                                if self.debug:
                                    print(f"无法设置删除线(方法2): {e}")
            
            return content
        except Exception as e:
            if self.debug:
                print(f"错误: 处理内联标签失败: {e}")
            # 简单处理，直接添加纯文本
            clean_text = re.sub(r'<[^>]*>', ' ', content)
            paragraph.add_run(clean_text.strip())
            return content
    
    def _fallback_convert(self, html_content: str) -> Paragraph:
        """基本HTML转换，处理常见的HTML标签
        
        Args:
            html_content: HTML内容
            
        Returns:
            Paragraph: 创建的段落
        """
        if self.debug:
            print("使用基本HTML转换")
        
        # 创建新段落
        paragraph = self.document.add_paragraph()
        
        # 简单处理一些基本HTML标签
        # 这里只是一个非常基础的实现，无法处理复杂的HTML
        content = html_content
        
        # 移除HTML标签（非常简单的实现）
        import re
        content = re.sub(r'<[^>]*>', ' ', content)
        
        # 添加文本
        paragraph.add_run(content)
        
        return paragraph 