"""
表格转换集成测试
"""
import pytest
from docx import Document

from src.converter import BaseConverter


class TestTableIntegration:
    """表格转换集成测试"""

    @pytest.fixture
    def base_converter(self):
        """创建基础转换器实例"""
        return BaseConverter(debug=False)

    def test_basic_table(self, base_converter):
        """测试基本表格转换"""
        md_text = """
| 姓名 | 年龄 | 职业 |
| ---- | ---- | ---- |
| 张三 | 25 | 工程师 |
| 李四 | 30 | 设计师 |
| 王五 | 28 | 产品经理 |
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        # 验证表格结构
        assert len(table.rows) >= 3  # 至少有3行
        assert len(table.columns) == 3
        
        # 验证表格内容 - 只验证存在性，不验证具体位置
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text.strip())
        
        assert "姓名" in cell_texts
        assert "年龄" in cell_texts
        assert "职业" in cell_texts
        assert "张三" in cell_texts
        assert "25" in cell_texts
        assert "工程师" in cell_texts
        assert "李四" in cell_texts
        assert "30" in cell_texts
        assert "设计师" in cell_texts

    def test_table_with_alignment(self, base_converter):
        """测试带对齐方式的表格"""
        md_text = """
| 左对齐 | 居中对齐 | 右对齐 |
| :---- | :----: | ----: |
| 内容 | 内容 | 内容 |
| 长一点的内容 | 长一点的内容 | 长一点的内容 |
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        # 验证表格结构
        assert len(table.rows) >= 3  # 至少有3行
        assert len(table.columns) == 3
        
        # 验证表格内容 - 只验证存在性，不验证具体位置
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text.strip())
        
        assert "左对齐" in cell_texts
        assert "居中对齐" in cell_texts
        assert "右对齐" in cell_texts
        assert "内容" in cell_texts
        assert "长一点的内容" in cell_texts

    def test_complex_table(self, base_converter):
        """测试复杂表格（包含样式）"""
        md_text = """
| 项目 | 描述 | 价格 | 数量 |
| ---- | ---- | ----: | :----: |
| 商品A | 这是一个**加粗**的描述 | 100.00 | 2 |
| 商品B | 这是一个*斜体*的描述 | 85.50 | 1 |
| 商品C | 这是一个~~删除线~~的描述 | 50.00 | 5 |
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        # 验证表格结构
        assert len(table.rows) >= 3  # 至少有3行
        assert len(table.columns) == 4
        
        # 验证表格内容 - 只验证存在性，不验证具体位置
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text.strip())
        
        assert "项目" in cell_texts
        assert "描述" in cell_texts
        assert "价格" in cell_texts
        assert "数量" in cell_texts
        assert "商品A" in cell_texts
        assert "加粗" in " ".join(cell_texts)  # 检查文本中是否包含"加粗"
        assert "100.00" in cell_texts
        assert "2" in cell_texts

    def test_empty_cells(self, base_converter):
        """测试空单元格"""
        md_text = """
| A | B | C |
| ---- | ---- | ---- |
| 数据 |  | 数据 |
|  | 数据 | 数据 |
| 数据 | 数据 |  |
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        # 验证表格结构
        assert len(table.rows) >= 4  # 至少有4行
        assert len(table.columns) == 3
        
        # 验证表格内容 - 只验证存在性，不验证具体位置
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text.strip())
        
        assert "A" in cell_texts
        assert "B" in cell_texts
        assert "C" in cell_texts
        assert "数据" in cell_texts
        assert "" in cell_texts  # 检查是否有空单元格

    def test_multiple_tables(self, base_converter):
        """测试多个表格"""
        md_text = """
# 第一个表格

| A | B |
| --- | --- |
| 1 | 2 |

# 第二个表格

| C | D |
| --- | --- |
| 3 | 4 |
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        assert len(doc.tables) == 2
        
        # 验证第一个表格
        table1 = doc.tables[0]
        assert len(table1.rows) >= 2  # 至少有2行
        assert len(table1.columns) == 2
        
        # 验证第一个表格内容 - 只验证存在性，不验证具体位置
        cell_texts1 = []
        for row in table1.rows:
            for cell in row.cells:
                cell_texts1.append(cell.text.strip())
        
        assert "A" in cell_texts1
        assert "B" in cell_texts1
        assert "1" in cell_texts1
        assert "2" in cell_texts1
        
        # 验证第二个表格
        table2 = doc.tables[1]
        assert len(table2.rows) >= 2  # 至少有2行
        assert len(table2.columns) == 2
        
        # 验证第二个表格内容 - 只验证存在性，不验证具体位置
        cell_texts2 = []
        for row in table2.rows:
            for cell in row.cells:
                cell_texts2.append(cell.text.strip())
        
        assert "C" in cell_texts2
        assert "D" in cell_texts2
        assert "3" in cell_texts2
        assert "4" in cell_texts2 