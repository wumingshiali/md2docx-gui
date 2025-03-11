"""
分隔线转换集成测试
"""
import pytest
from docx import Document

from src.converter import BaseConverter


class TestHRIntegration:
    """分隔线转换集成测试"""

    @pytest.fixture
    def base_converter(self):
        """创建基础转换器实例"""
        return BaseConverter(debug=False)

    def test_basic_hr(self, base_converter):
        """测试基本分隔线转换"""
        md_text = """
这是一段文本

---

这是另一段文本
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该有3个段落：文本、分隔线、文本
        assert len(doc.paragraphs) >= 3
        
        # 验证分隔线段落的对齐方式
        # 注意：分隔线是第二个段落
        hr_paragraph = doc.paragraphs[1]
        assert hr_paragraph.alignment == 1  # 居中对齐

    def test_multiple_hr(self, base_converter):
        """测试多个分隔线"""
        md_text = """
第一段

---

第二段

***

第三段

___

第四段
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该有7个段落：文本、分隔线、文本、分隔线、文本、分隔线、文本
        assert len(doc.paragraphs) >= 7
        
        # 验证分隔线段落的对齐方式
        hr_paragraphs = [doc.paragraphs[1], doc.paragraphs[3], doc.paragraphs[5]]
        for hr_paragraph in hr_paragraphs:
            assert hr_paragraph.alignment == 1  # 居中对齐

    def test_hr_with_other_elements(self, base_converter):
        """测试分隔线与其他元素混合"""
        md_text = """
# 标题

这是一段文本

---

- 列表项1
- 列表项2

---

> 引用内容
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该至少有6个段落：标题、文本、分隔线、列表项1、列表项2、分隔线、引用
        assert len(doc.paragraphs) >= 6
        
        # 验证分隔线段落的对齐方式
        # 注意：分隔线是第3个和第6个段落
        hr_paragraphs = [doc.paragraphs[2], doc.paragraphs[5]]
        for hr_paragraph in hr_paragraphs:
            assert hr_paragraph.alignment == 1  # 居中对齐 