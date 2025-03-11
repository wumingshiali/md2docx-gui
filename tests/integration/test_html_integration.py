"""
HTML转换集成测试
"""
import pytest
from docx import Document

from src.converter import BaseConverter
from src.converter.elements.html import HTML2DOCX_AVAILABLE


class TestHtmlIntegration:
    """HTML转换集成测试"""

    @pytest.fixture
    def base_converter(self):
        """创建基础转换器实例"""
        return BaseConverter(debug=False)

    def test_basic_html(self, base_converter):
        """测试基本HTML转换"""
        md_text = """
这是普通文本。

<div>这是一个HTML div</div>

这是更多普通文本。
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该有至少3个段落
        assert len(doc.paragraphs) >= 3
        
        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "这是普通文本" in all_text
        assert "这是一个HTML div" in all_text
        assert "这是更多普通文本" in all_text

    def test_html_with_formatting(self, base_converter):
        """测试带格式的HTML转换"""
        md_text = """
<p>这是一个<strong>粗体</strong>和<em>斜体</em>文本。</p>
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该有至少1个段落
        assert len(doc.paragraphs) >= 1
        
        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "这是一个" in all_text
        assert "粗体" in all_text
        assert "斜体" in all_text
        assert "文本" in all_text

    def test_mixed_markdown_and_html(self, base_converter):
        """测试混合Markdown和HTML"""
        md_text = """
# 标题

这是**Markdown**格式的文本。

<div>
  <p>这是<strong>HTML</strong>格式的文本。</p>
</div>

1. 列表项1
2. 列表项2
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该有多个段落
        assert len(doc.paragraphs) >= 5
        
        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "标题" in all_text
        assert "Markdown" in all_text
        assert "HTML" in all_text
        assert "列表项1" in all_text
        assert "列表项2" in all_text

    def test_html_table(self, base_converter):
        """测试HTML表格"""
        md_text = """
<table>
  <tr>
    <th>姓名</th>
    <th>年龄</th>
  </tr>
  <tr>
    <td>张三</td>
    <td>25</td>
  </tr>
</table>
"""
        # 转换Markdown到DOCX
        doc = base_converter.convert(md_text)
        
        # 确保文档不为空
        assert len(doc.paragraphs) > 0
        
        # 检查表格是否存在
        assert len(doc.tables) > 0, "文档中应该包含表格"
        
        # 检查表格内容
        if len(doc.tables) > 0:
            table = doc.tables[0]
            # 获取表格中的所有文本
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
            
            # 检查表格内容是否包含预期的文本
            assert "姓名" in table_text or "年龄" in table_text or "张三" in table_text or "25" in table_text, f"表格内容: {table_text}"

    @pytest.mark.skipif(not HTML2DOCX_AVAILABLE, reason="html2docx not available")
    def test_complex_html(self, base_converter):
        """测试复杂HTML结构"""
        md_text = """
<div class="container">
  <h2>HTML标题</h2>
  <p>这是一个<strong>复杂</strong>的<em>HTML</em>结构。</p>
  <ul>
    <li>项目1</li>
    <li>项目2</li>
    <li>项目3</li>
  </ul>
</div>
"""
        doc = base_converter.convert(md_text)
        
        # 验证结果
        # 应该有多个段落
        assert len(doc.paragraphs) > 0
        
        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "HTML标题" in all_text
        assert "复杂" in all_text
        assert "HTML" in all_text
        assert "项目" in all_text 