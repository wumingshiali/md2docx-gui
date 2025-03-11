"""
引用块转换器测试
"""
import pytest
from docx import Document
from markdown_it import MarkdownIt
from src.converter.base import BaseConverter
from src.converter.elements import BlockquoteConverter, TextConverter


def test_basic_blockquote():
    """测试基本引用块"""
    md_text = "> 这是一个简单的引用。"
    
    converter = BaseConverter()
    converter.register_converter('blockquote', BlockquoteConverter())
    converter.register_converter('paragraph', TextConverter())
    
    doc = converter.convert(md_text)
    assert len(doc.paragraphs) == 1
    # 引用块应该有特殊的样式
    paragraph = doc.paragraphs[0]
    assert paragraph.text == "这是一个简单的引用。"
    assert paragraph.style.name == "Quote"


def test_multi_line_blockquote():
    """测试多行引用块"""
    md_text = "> 第一行引用\n> 第二行引用\n> 第三行引用"
    
    converter = BaseConverter()
    converter.register_converter('blockquote', BlockquoteConverter())
    converter.register_converter('paragraph', TextConverter())
    
    doc = converter.convert(md_text)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "第一行引用 第二行引用 第三行引用"


def test_nested_blockquote():
    """测试嵌套引用块"""
    md_text = "> 外层引用\n>> 内层引用"
    
    converter = BaseConverter()
    converter.register_converter('blockquote', BlockquoteConverter())
    converter.register_converter('paragraph', TextConverter())
    
    doc = converter.convert(md_text)
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].text == "外层引用"
    assert doc.paragraphs[1].text == "内层引用"
    # 内层引用应该有不同的缩进
    assert doc.paragraphs[1].style.name == "Quote2"


def test_blockquote_with_styles():
    """测试带样式的引用块"""
    md_text = "> **粗体**和*斜体*文本"
    
    converter = BaseConverter()
    converter.register_converter('blockquote', BlockquoteConverter())
    converter.register_converter('paragraph', TextConverter())
    
    doc = converter.convert(md_text)
    runs = doc.paragraphs[0].runs
    
    assert runs[0].text == "粗体" and runs[0].bold
    assert runs[1].text == "和"
    assert runs[2].text == "斜体" and runs[2].italic
    assert runs[3].text == "文本"


def test_empty_blockquote():
    """测试空引用块"""
    md_text = ">"
    
    converter = BaseConverter()
    converter.register_converter('blockquote', BlockquoteConverter())
    converter.register_converter('paragraph', TextConverter())
    
    doc = converter.convert(md_text)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "" 