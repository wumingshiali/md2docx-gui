"""
表格转换器单元测试
"""
import pytest
from docx import Document
from unittest.mock import MagicMock, patch

from src.converter.elements.table import TableConverter
from src.converter import BaseConverter


def test_init():
    """测试初始化"""
    converter = TableConverter()
    assert converter is not None
    assert converter.document is None
    assert converter.base_converter is None
    assert converter.debug is False
    
    # 测试带基础转换器的初始化
    base_converter = MagicMock()
    base_converter.debug = True
    converter = TableConverter(base_converter)
    assert converter.base_converter == base_converter
    assert converter.debug is True


def test_document_not_set():
    """测试文档未设置的情况"""
    converter = TableConverter()
    with pytest.raises(ValueError):
        converter.convert(MagicMock())


def test_basic_table():
    """测试基本表格转换"""
    # 创建转换器
    converter = TableConverter()
    converter.set_document(Document())
    
    # 创建模拟表格token
    table_token = MagicMock()
    table_token.type = 'table_open'
    
    # 创建表头行
    header_row = MagicMock()
    header_row.type = 'tr'
    header_cell1 = MagicMock()
    header_cell1.type = 'th'
    header_cell1.children = [MagicMock(type='text', content='标题1')]
    header_cell1.attrs = {}
    
    header_cell2 = MagicMock()
    header_cell2.type = 'th'
    header_cell2.children = [MagicMock(type='text', content='标题2')]
    header_cell2.attrs = {}
    
    header_row.children = [header_cell1, header_cell2]
    
    # 创建数据行
    data_row = MagicMock()
    data_row.type = 'tr'
    data_cell1 = MagicMock()
    data_cell1.type = 'td'
    data_cell1.children = [MagicMock(type='text', content='数据1')]
    data_cell1.attrs = {}
    
    data_cell2 = MagicMock()
    data_cell2.type = 'td'
    data_cell2.children = [MagicMock(type='text', content='数据2')]
    data_cell2.attrs = {}
    
    data_row.children = [data_cell1, data_cell2]
    
    # 设置表格的行
    table_token.children = [header_row, data_row]
    
    # 转换表格
    table = converter.convert(table_token)
    
    # 验证结果
    assert table is not None
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    
    # 由于我们没有使用基础转换器，所以这里只能验证表格结构
    # 实际内容需要在集成测试中验证


def test_table_with_alignment():
    """测试带对齐方式的表格"""
    # 创建转换器
    converter = TableConverter()
    converter.set_document(Document())
    
    # 创建模拟表格token
    table_token = MagicMock()
    table_token.type = 'table_open'
    
    # 创建表头行
    header_row = MagicMock()
    header_row.type = 'tr'
    header_cell1 = MagicMock()
    header_cell1.type = 'th'
    header_cell1.children = [MagicMock(type='text', content='左对齐')]
    header_cell1.attrs = {'align': 'left'}
    
    header_cell2 = MagicMock()
    header_cell2.type = 'th'
    header_cell2.children = [MagicMock(type='text', content='居中')]
    header_cell2.attrs = {'align': 'center'}
    
    header_cell3 = MagicMock()
    header_cell3.type = 'th'
    header_cell3.children = [MagicMock(type='text', content='右对齐')]
    header_cell3.attrs = {'align': 'right'}
    
    header_row.children = [header_cell1, header_cell2, header_cell3]
    
    # 创建数据行
    data_row = MagicMock()
    data_row.type = 'tr'
    data_cell1 = MagicMock()
    data_cell1.type = 'td'
    data_cell1.children = [MagicMock(type='text', content='数据1')]
    data_cell1.attrs = {'align': 'left'}
    
    data_cell2 = MagicMock()
    data_cell2.type = 'td'
    data_cell2.children = [MagicMock(type='text', content='数据2')]
    data_cell2.attrs = {'align': 'center'}
    
    data_cell3 = MagicMock()
    data_cell3.type = 'td'
    data_cell3.children = [MagicMock(type='text', content='数据3')]
    data_cell3.attrs = {'align': 'right'}
    
    data_row.children = [data_cell1, data_cell2, data_cell3]
    
    # 设置表格的行
    table_token.children = [header_row, data_row]
    
    # 转换表格
    table = converter.convert(table_token)
    
    # 验证结果
    assert table is not None
    assert len(table.rows) == 2
    assert len(table.columns) == 3


def test_empty_table():
    """测试空表格"""
    # 创建转换器
    converter = TableConverter()
    converter.set_document(Document())
    
    # 创建模拟表格token
    table_token = MagicMock()
    table_token.type = 'table_open'
    table_token.children = []
    
    # 转换表格
    table = converter.convert(table_token)
    
    # 验证结果
    assert table is None


def test_table_with_base_converter():
    """测试使用基础转换器处理表格内容"""
    # 创建基础转换器
    base_converter = MagicMock()
    base_converter._process_inline = MagicMock()
    base_converter._process_token = MagicMock()
    
    # 创建转换器
    converter = TableConverter(base_converter)
    converter.set_document(Document())
    
    # 创建模拟表格token
    table_token = MagicMock()
    table_token.type = 'table_open'
    
    # 创建表头行
    header_row = MagicMock()
    header_row.type = 'tr'
    header_cell = MagicMock()
    header_cell.type = 'th'
    inline_token = MagicMock(type='inline')
    header_cell.children = [inline_token]
    header_cell.attrs = {}
    header_row.children = [header_cell]
    
    # 创建数据行
    data_row = MagicMock()
    data_row.type = 'tr'
    data_cell = MagicMock()
    data_cell.type = 'td'
    data_cell.children = [inline_token]
    data_cell.attrs = {}
    data_row.children = [data_cell]
    
    # 设置表格的行
    table_token.children = [header_row, data_row]
    
    # 转换表格
    table = converter.convert(table_token)
    
    # 验证结果
    assert table is not None
    assert len(table.rows) == 2
    assert len(table.columns) == 1
    
    # 验证基础转换器被调用
    assert base_converter._process_inline.call_count == 2 