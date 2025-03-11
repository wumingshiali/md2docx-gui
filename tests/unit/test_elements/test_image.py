"""
测试图片转换功能
"""
import os
import pytest
from unittest.mock import MagicMock, patch
from io import BytesIO
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_ALIGN_PARAGRAPH

from markdown_it import MarkdownIt
from src.converter.elements.image import ImageConverter


class TestImageConverter:
    """测试图片转换器"""

    @pytest.fixture
    def image_converter(self):
        """创建图片转换器实例"""
        base_converter = MagicMock()
        base_converter.debug = False
        converter = ImageConverter(base_converter)
        document = Document()
        converter.set_document(document)
        return converter

    @pytest.fixture
    def md_parser(self):
        """创建Markdown解析器"""
        return MarkdownIt()

    def test_init(self, image_converter):
        """测试初始化"""
        assert image_converter.document is not None
        assert image_converter._image_cache == {}

    def test_parse_size(self, image_converter):
        """测试尺寸解析"""
        # 测试无尺寸
        assert image_converter._parse_size("图片") == (None, None)
        
        # 测试有尺寸
        assert image_converter._parse_size("图片|100x200") == (100, 200)
        
        # 测试格式错误
        assert image_converter._parse_size("图片|100") == (None, None)
        assert image_converter._parse_size("图片|axb") == (None, None)
        assert image_converter._parse_size("图片|100x") == (None, None)

    @patch('src.converter.elements.image.requests.get')
    def test_get_image_data_online(self, mock_get, image_converter):
        """测试获取在线图片数据"""
        # 模拟请求响应
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.content = b'fake_image_data'
        mock_get.return_value = mock_response
        
        # 测试获取在线图片
        url = 'http://example.com/image.png'
        result = image_converter._get_image_data(url)
        
        # 验证结果
        assert isinstance(result, BytesIO)
        assert result.getvalue() == b'fake_image_data'
        
        # 验证缓存
        assert url in image_converter._image_cache
        assert image_converter._image_cache[url] == b'fake_image_data'
        
        # 再次获取应该使用缓存
        mock_get.reset_mock()
        result2 = image_converter._get_image_data(url)
        assert isinstance(result2, BytesIO)
        assert result2.getvalue() == b'fake_image_data'
        mock_get.assert_not_called()

    @patch('src.converter.elements.image.os.path.exists')
    @patch('builtins.open')
    def test_get_image_data_local(self, mock_open, mock_exists, image_converter):
        """测试获取本地图片数据"""
        # 模拟文件存在
        mock_exists.return_value = True
        
        # 模拟文件读取
        mock_file = MagicMock()
        mock_file.__enter__.return_value.read.return_value = b'fake_local_image'
        mock_open.return_value = mock_file
        
        # 测试获取本地图片
        path = 'image.png'
        result = image_converter._get_image_data(path)
        
        # 验证结果
        assert isinstance(result, BytesIO)
        assert result.getvalue() == b'fake_local_image'
        
        # 验证缓存
        assert path in image_converter._image_cache
        assert image_converter._image_cache[path] == b'fake_local_image'

    def test_convert_in_paragraph(self, image_converter):
        """测试在段落中转换图片"""
        # 创建段落
        paragraph = image_converter.document.add_paragraph()
        
        # 创建图片标记
        token = MagicMock()
        token.attrs = {'src': 'test.png'}
        token.content = '测试图片'
        
        # 模拟获取图片数据
        with patch.object(image_converter, '_get_image_data') as mock_get_data:
            mock_get_data.return_value = BytesIO(b'fake_image_data')
            
            # 测试转换
            image_converter.convert_in_paragraph(paragraph, token)
            
            # 验证结果
            mock_get_data.assert_called_once_with('test.png')
            assert len(paragraph.runs) > 0

    def test_convert_in_paragraph_with_size(self, image_converter):
        """测试在段落中转换带尺寸的图片"""
        # 创建段落
        paragraph = image_converter.document.add_paragraph()
        
        # 创建图片标记
        token = MagicMock()
        token.attrs = {'src': 'test.png'}
        token.content = '测试图片|150x100'
        
        # 模拟获取图片数据和解析尺寸
        with patch.object(image_converter, '_get_image_data') as mock_get_data, \
             patch.object(image_converter, '_parse_size') as mock_parse_size:
            mock_get_data.return_value = BytesIO(b'fake_image_data')
            mock_parse_size.return_value = (150, 100)
            
            # 测试转换
            image_converter.convert_in_paragraph(paragraph, token)
            
            # 验证结果
            mock_get_data.assert_called_once_with('test.png')
            mock_parse_size.assert_called_once_with('测试图片|150x100')
            assert len(paragraph.runs) > 0

    def test_convert(self, image_converter):
        """测试转换图片元素"""
        # 创建图片标记
        token = MagicMock()
        token.attrs = {'src': 'test.png', 'title': '图片标题'}
        token.content = '测试图片'
        
        # 模拟获取图片数据
        with patch.object(image_converter, '_get_image_data') as mock_get_data:
            mock_get_data.return_value = BytesIO(b'fake_image_data')
            
            # 测试转换
            image_converter.convert((token, token))
            
            # 验证结果
            mock_get_data.assert_called_once_with('test.png')
            
            # 验证段落数量
            # 注意：根据实际实现，可能是1个段落（图片和标题在同一段落）或2个段落（图片和标题分开）
            assert len(image_converter.document.paragraphs) >= 1
            
            # 验证图片段落
            assert len(image_converter.document.paragraphs[0].runs) > 0
            
            # 如果有标题段落，验证标题内容
            if len(image_converter.document.paragraphs) > 1:
                assert image_converter.document.paragraphs[1].runs[0].text == '图片标题'
                assert image_converter.document.paragraphs[1].runs[0].italic

    def test_convert_with_size(self, image_converter):
        """测试转换带尺寸的图片元素"""
        # 创建图片标记
        token = MagicMock()
        token.attrs = {'src': 'test.png'}
        token.content = '测试图片|200x150'
        
        # 模拟获取图片数据和解析尺寸
        with patch.object(image_converter, '_get_image_data') as mock_get_data, \
             patch.object(image_converter, '_parse_size') as mock_parse_size:
            mock_get_data.return_value = BytesIO(b'fake_image_data')
            mock_parse_size.return_value = (200, 150)
            
            # 测试转换
            image_converter.convert((token, token))
            
            # 验证结果
            mock_get_data.assert_called_once_with('test.png')
            mock_parse_size.assert_called_once_with('测试图片|200x150')
            assert len(image_converter.document.paragraphs) == 1
            assert image_converter.document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
            assert len(image_converter.document.paragraphs[0].runs) > 0

    def test_integration_basic_image(self, md_parser):
        """集成测试：基本图片"""
        md_text = "![测试图片](test.png)"
        tokens = md_parser.parse(md_text)
        
        # 验证解析结果
        assert len(tokens) == 3  # paragraph_open, inline, paragraph_close
        assert tokens[1].children[0].type == 'image'
        assert tokens[1].children[0].attrs['src'] == 'test.png'
        assert tokens[1].children[0].content == '测试图片'

    def test_integration_image_with_title(self, md_parser):
        """集成测试：带标题的图片"""
        md_text = '![测试图片](test.png "图片标题")'
        tokens = md_parser.parse(md_text)
        
        # 验证解析结果
        assert tokens[1].children[0].type == 'image'
        assert tokens[1].children[0].attrs['src'] == 'test.png'
        assert tokens[1].children[0].attrs['title'] == '图片标题'
        assert tokens[1].children[0].content == '测试图片'

    def test_integration_image_with_size(self, md_parser):
        """集成测试：带尺寸的图片"""
        md_text = "![测试图片|200x150](test.png)"
        tokens = md_parser.parse(md_text)
        
        # 验证解析结果
        assert tokens[1].children[0].type == 'image'
        assert tokens[1].children[0].attrs['src'] == 'test.png'
        assert tokens[1].children[0].content == '测试图片|200x150'

    def test_integration_inline_image(self, md_parser):
        """集成测试：内联图片"""
        md_text = "这是一个段落，包含一个内联图片 ![内联图片](test.png) 在文本中。"
        tokens = md_parser.parse(md_text)
        
        # 验证解析结果
        assert tokens[1].children[1].type == 'image'
        assert tokens[1].children[1].attrs['src'] == 'test.png'
        assert tokens[1].children[1].content == '内联图片' 