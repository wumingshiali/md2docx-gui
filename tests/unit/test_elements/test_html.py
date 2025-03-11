"""
HTML转换器单元测试
"""
import pytest
from docx import Document
from unittest.mock import MagicMock, patch

from src.converter.elements.html import HtmlConverter, HTML2DOCX_AVAILABLE


def test_init():
    """测试初始化"""
    converter = HtmlConverter()
    assert converter is not None
    assert converter.document is None
    assert converter.debug is False
    
    # 测试带基础转换器的初始化
    base_converter = MagicMock()
    base_converter.debug = True
    converter = HtmlConverter(base_converter)
    assert converter.debug is True


def test_document_not_set():
    """测试文档未设置的情况"""
    converter = HtmlConverter()
    with pytest.raises(ValueError):
        converter.convert(MagicMock())


@pytest.mark.skipif(not HTML2DOCX_AVAILABLE, reason="html2docx not available")
def test_convert_with_html2docx():
    """测试使用html2docx转换HTML"""
    # 创建转换器
    converter = HtmlConverter()
    converter.set_document(Document())
    
    # 创建模拟HTML token
    token = MagicMock()
    token.type = 'html_block'
    token.content = '<p>这是一个<strong>HTML</strong>段落</p>'
    
    # 转换HTML
    result = converter.convert(token)
    
    # 验证结果
    assert result is not None
    # 由于html2docx的具体行为难以模拟，这里只验证基本结果


def test_convert_without_html2docx():
    """测试在html2docx不可用时的转换"""
    # 创建转换器
    converter = HtmlConverter()
    converter.set_document(Document())
    
    # 模拟html2docx不可用
    with patch('src.converter.elements.html.HTML2DOCX_AVAILABLE', False):
        # 创建模拟HTML token
        token = MagicMock()
        token.type = 'html_block'
        token.content = '<p>这是一个<strong>HTML</strong>段落</p>'
        
        # 转换HTML
        result = converter.convert(token)
        
        # 验证结果
        assert result is not None
        assert len(converter.document.paragraphs) > 0


def test_convert_empty_html():
    """测试转换空HTML内容"""
    # 创建转换器
    converter = HtmlConverter()
    converter.set_document(Document())
    
    # 创建模拟空HTML token
    token = MagicMock()
    token.type = 'html_block'
    token.content = ''
    
    # 转换HTML
    result = converter.convert(token)
    
    # 验证结果
    assert result is None


def test_fallback_convert():
    """测试回退转换方法"""
    # 创建转换器
    converter = HtmlConverter()
    converter.set_document(Document())
    
    # 调用回退转换方法
    html_content = '<p>这是一个<strong>HTML</strong>段落</p>'
    result = converter._fallback_convert(html_content)
    
    # 验证结果
    assert result is not None
    assert len(converter.document.paragraphs) > 0
    # 验证HTML标签被移除
    assert '<' not in result.text
    assert '>' not in result.text 