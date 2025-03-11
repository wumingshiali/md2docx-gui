import pytest
from docx.shared import RGBColor

from src.converter.elements.code import CodeConverter
from src.converter.base import BaseConverter


def test_basic_code_block(base_converter):
    """测试基本代码块的转换"""
    markdown = """```python
def hello():
    print("Hello, World!")
```"""
    doc = base_converter.convert(markdown)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 1
    assert paragraphs[0].text == 'def hello():\n    print("Hello, World!")'


def test_code_block_with_language(base_converter):
    """测试带有语言标识的代码块"""
    markdown = """```javascript
function hello() {
    console.log("Hello, World!");
}
```"""
    doc = base_converter.convert(markdown)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 1
    assert paragraphs[0].text == 'function hello() {\n    console.log("Hello, World!");\n}'


def test_code_block_without_language(base_converter):
    """测试不带语言标识的代码块"""
    markdown = """```
Some code without language
```"""
    doc = base_converter.convert(markdown)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 1
    assert paragraphs[0].text == 'Some code without language'


def test_multiple_code_blocks(base_converter):
    """测试多个代码块"""
    markdown = """```python
def hello():
    print("Hello")
```

```javascript
console.log("World");
```"""
    doc = base_converter.convert(markdown)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 3  # 两个代码块加一个空行
    assert paragraphs[0].text == 'def hello():\n    print("Hello")'
    assert paragraphs[2].text == 'console.log("World");'


def test_empty_code_block(base_converter):
    """测试空代码块"""
    markdown = """```python
```"""
    doc = base_converter.convert(markdown)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 1
    assert paragraphs[0].text == ''


def test_code_block_with_special_characters(base_converter):
    """测试包含特殊字符的代码块"""
    markdown = """```python
def special_chars():
    # 这是一个注释
    print("特殊字符：!@#$%^&*()")
```"""
    doc = base_converter.convert(markdown)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 1
    assert paragraphs[0].text == 'def special_chars():\n    # 这是一个注释\n    print("特殊字符：!@#$%^&*()")' 