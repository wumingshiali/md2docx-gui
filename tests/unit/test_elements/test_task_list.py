"""
任务列表转换器单元测试
"""
import pytest
from docx import Document
from unittest.mock import MagicMock, patch

from src.converter.elements.task_list import TaskListConverter


def test_init():
    """测试初始化"""
    converter = TaskListConverter()
    assert converter is not None
    assert converter.document is None
    assert converter.debug is False
    assert converter.list_converter is None
    
    # 测试带基础转换器的初始化
    base_converter = MagicMock()
    base_converter.debug = True
    base_converter.converters = {'list': MagicMock()}
    converter = TaskListConverter(base_converter)
    assert converter.debug is True
    assert converter.list_converter is not None


def test_document_not_set():
    """测试文档未设置的情况"""
    converter = TaskListConverter()
    with pytest.raises(ValueError):
        converter.convert((MagicMock(), MagicMock()))


def test_convert_checked_task():
    """测试已选中的任务列表项转换"""
    # 创建转换器
    converter = TaskListConverter()
    converter.set_document(Document())
    
    # 模拟列表转换器
    list_converter = MagicMock()
    paragraph = Document().add_paragraph("任务项")
    list_converter.convert.return_value = paragraph
    converter.list_converter = list_converter
    
    # 创建模拟任务列表token
    list_token = MagicMock()
    list_token.type = 'bullet_list_open'
    
    # 创建模拟内容token - 使用字符串内容而不是复杂的模拟对象
    content_token = MagicMock()
    content_token.type = 'inline'
    content_token.content = "[x] 已完成任务"
    content_token.children = []
    
    # 转换任务列表项
    paragraph = converter.convert((list_token, content_token))
    
    # 验证结果
    assert paragraph is not None
    assert list_converter.convert.called
    assert paragraph._element is not None


def test_convert_unchecked_task():
    """测试未选中的任务列表项转换"""
    # 创建转换器
    converter = TaskListConverter()
    converter.set_document(Document())
    
    # 模拟列表转换器
    list_converter = MagicMock()
    paragraph = Document().add_paragraph("任务项")
    list_converter.convert.return_value = paragraph
    converter.list_converter = list_converter
    
    # 创建模拟任务列表token
    list_token = MagicMock()
    list_token.type = 'bullet_list_open'
    
    # 创建模拟内容token - 使用字符串内容而不是复杂的模拟对象
    content_token = MagicMock()
    content_token.type = 'inline'
    content_token.content = "[ ] 未完成任务"
    content_token.children = []
    
    # 转换任务列表项
    paragraph = converter.convert((list_token, content_token))
    
    # 验证结果
    assert paragraph is not None
    assert list_converter.convert.called
    assert paragraph._element is not None


def test_convert_without_list_converter():
    """测试没有列表转换器的情况"""
    # 创建转换器
    converter = TaskListConverter()
    converter.set_document(Document())
    converter.list_converter = None
    
    # 创建模拟任务列表token
    list_token = MagicMock()
    list_token.type = 'bullet_list_open'
    
    # 创建模拟内容token - 使用字符串内容而不是复杂的模拟对象
    content_token = MagicMock()
    content_token.type = 'inline'
    content_token.content = "[x] 任务内容"
    content_token.children = []
    
    # 转换任务列表项
    paragraph = converter.convert((list_token, content_token))
    
    # 验证结果
    assert paragraph is not None
    assert paragraph._element is not None 